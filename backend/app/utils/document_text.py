from __future__ import annotations

import io

# Extract plain text from an uploaded contract file. Supports .txt, .pdf (pypdf),
# and .docx (python-docx). Lazy imports so the app still boots if a parser isn't
# installed — the caller surfaces a friendly message instead.


class UnsupportedDocument(Exception):
    pass


def _from_pdf(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:  # pragma: no cover
        raise UnsupportedDocument("PDF support isn't installed on the server. Paste the contract text instead, " "or install pypdf.")
    reader = PdfReader(io.BytesIO(raw))
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(parts)


def _from_docx(raw: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError:  # pragma: no cover
        raise UnsupportedDocument("Word (.docx) support isn't installed on the server. Paste the contract text " "instead, or install python-docx.")
    document = docx.Document(io.BytesIO(raw))
    return "\n".join(p.text for p in document.paragraphs)


def extract_text(filename: str, content_type: str | None, raw: bytes) -> str:
    """Return the document's text, or raise UnsupportedDocument."""
    name = (filename or "").lower()
    ctype = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ctype:
        text = _from_pdf(raw)
    elif name.endswith(".docx") or "word" in ctype or "officedocument" in ctype:
        text = _from_docx(raw)
    elif name.endswith((".txt", ".md", ".text")) or ctype.startswith("text/"):
        text = raw.decode("utf-8", errors="replace")
    else:
        # Last resort: try UTF-8; if it's binary garbage the reviewer will catch it.
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            raise UnsupportedDocument("Unsupported file type. Upload a PDF, Word (.docx) or text file, or paste the text in instead.")

    text = text.replace("\x00", "").strip()
    if len(text) < 20:
        raise UnsupportedDocument("Couldn't extract readable text from this file (it may be a scanned image). " "Paste the contract text instead.")
    return text
